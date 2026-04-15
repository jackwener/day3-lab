# -*- coding: utf-8 -*-
"""报告导出服务"""

import os
from datetime import datetime
from io import BytesIO


class ExportService:
    """报告导出服务"""
    
    def export_pdf(self, review, issues):
        """导出PDF格式审核报告"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4)
            
            # 注册中文字体（如果有的话）
            try:
                pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
                font_name = 'SimSun'
            except:
                font_name = 'Helvetica'
            
            styles = getSampleStyleSheet()
            
            # 标题样式
            title_style = ParagraphStyle(
                'Title',
                parent=styles['Title'],
                fontName=font_name,
                fontSize=18,
                spaceAfter=30
            )
            
            # 正文样式
            body_style = ParagraphStyle(
                'Body',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=10,
                spaceAfter=10
            )
            
            elements = []
            
            # 标题
            elements.append(Paragraph(f"审核报告 - {review.id}", title_style))
            elements.append(Spacer(1, 0.5*cm))
            
            # 基本信息
            info_data = [
                ['研报标题', review.title or '未知'],
                ['作者', review.author or '未知'],
                ['研报类型', review.report_type],
                ['审核模式', review.mode],
                ['审核状态', review.status],
                ['综合评分', str(review.score) if review.score else '-'],
                ['提交时间', review.submitted_at.strftime('%Y-%m-%d %H:%M') if review.submitted_at else '-'],
                ['完成时间', review.completed_at.strftime('%Y-%m-%d %H:%M') if review.completed_at else '-'],
            ]
            
            info_table = Table(info_data, colWidths=[3*cm, 10*cm])
            info_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
            ]))
            elements.append(info_table)
            elements.append(Spacer(1, 1*cm))
            
            # 问题统计
            elements.append(Paragraph("问题统计", body_style))
            stats_data = [
                ['合规问题', str(review.compliance_issues)],
                ['内容问题', str(review.content_issues)],
                ['问题总数', str(review.compliance_issues + review.content_issues)],
            ]
            stats_table = Table(stats_data, colWidths=[5*cm, 3*cm])
            stats_table.setStyle(TableStyle([
                ('FONTNAME', (0, 0), (-1, -1), font_name),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            elements.append(stats_table)
            elements.append(Spacer(1, 1*cm))
            
            # 问题详情
            if issues:
                elements.append(Paragraph("问题详情", body_style))
                for i, issue in enumerate(issues, 1):
                    issue_text = f"{i}. [{issue.severity}] {issue.rule_name}"
                    elements.append(Paragraph(issue_text, body_style))
                    
                    if issue.location:
                        elements.append(Paragraph(f"   位置: {issue.location}", body_style))
                    if issue.excerpt:
                        elements.append(Paragraph(f"   内容: {issue.excerpt[:100]}...", body_style))
                    if issue.suggestion:
                        elements.append(Paragraph(f"   建议: {issue.suggestion}", body_style))
                    
                    elements.append(Spacer(1, 0.3*cm))
            
            doc.build(elements)
            buffer.seek(0)
            return buffer
            
        except ImportError:
            # 如果 reportlab 不可用，返回一个简单的文本
            return self._export_text(review, issues)
    
    def export_docx(self, review, issues):
        """导出DOCX格式审核报告"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # 标题
            title = doc.add_heading(f'审核报告 - {review.id}', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 基本信息
            doc.add_heading('基本信息', level=1)
            
            info_table = doc.add_table(rows=8, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ('研报标题', review.title or '未知'),
                ('作者', review.author or '未知'),
                ('研报类型', review.report_type),
                ('审核模式', review.mode),
                ('审核状态', review.status),
                ('综合评分', str(review.score) if review.score else '-'),
                ('提交时间', review.submitted_at.strftime('%Y-%m-%d %H:%M') if review.submitted_at else '-'),
                ('完成时间', review.completed_at.strftime('%Y-%m-%d %H:%M') if review.completed_at else '-'),
            ]
            
            for i, (label, value) in enumerate(info_data):
                row = info_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
            
            # 问题统计
            doc.add_heading('问题统计', level=1)
            doc.add_paragraph(f'合规问题: {review.compliance_issues}')
            doc.add_paragraph(f'内容问题: {review.content_issues}')
            doc.add_paragraph(f'问题总数: {review.compliance_issues + review.content_issues}')
            
            # 问题详情
            if issues:
                doc.add_heading('问题详情', level=1)
                
                for i, issue in enumerate(issues, 1):
                    doc.add_paragraph(f'{i}. [{issue.severity}] {issue.rule_name}')
                    
                    if issue.location:
                        doc.add_paragraph(f'   位置: {issue.location}')
                    if issue.excerpt:
                        doc.add_paragraph(f'   内容: {issue.excerpt[:200]}...' if len(issue.excerpt) > 200 else f'   内容: {issue.excerpt}')
                    if issue.suggestion:
                        doc.add_paragraph(f'   建议: {issue.suggestion}')
                    
                    doc.add_paragraph('')  # 空行
            
            # 保存到BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
            
        except ImportError:
            return self._export_text(review, issues)
    
    def _export_text(self, review, issues):
        """导出纯文本格式（备用）"""
        content = f"""
审核报告 - {review.id}

=== 基本信息 ===
研报标题: {review.title or '未知'}
作者: {review.author or '未知'}
研报类型: {review.report_type}
审核模式: {review.mode}
审核状态: {review.status}
综合评分: {review.score if review.score else '-'}
提交时间: {review.submitted_at.strftime('%Y-%m-%d %H:%M') if review.submitted_at else '-'}
完成时间: {review.completed_at.strftime('%Y-%m-%d %H:%M') if review.completed_at else '-'}

=== 问题统计 ===
合规问题: {review.compliance_issues}
内容问题: {review.content_issues}
问题总数: {review.compliance_issues + review.content_issues}

=== 问题详情 ===
"""
        for i, issue in enumerate(issues, 1):
            content += f"\n{i}. [{issue.severity}] {issue.rule_name}\n"
            if issue.location:
                content += f"   位置: {issue.location}\n"
            if issue.excerpt:
                content += f"   内容: {issue.excerpt[:200]}\n"
            if issue.suggestion:
                content += f"   建议: {issue.suggestion}\n"
        
        buffer = BytesIO(content.encode('utf-8'))
        buffer.seek(0)
        return buffer


# 单例
export_service = ExportService()
